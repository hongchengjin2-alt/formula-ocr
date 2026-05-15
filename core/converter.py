"""LaTeX to MathML/OMML conversion and LaTeX preview rendering."""

from io import BytesIO
from pathlib import Path
import re

import latex2mathml.converter
from lxml import etree
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PIL import Image

from core.paths import resource_path

# Path to MML2OMML.XSL (ships with this project, copied from MS Office)
_XSL_PATH = resource_path("assets", "MML2OMML.XSL")
_xslt = None
_MATHTEXT_SIZE_COMMAND_RE = re.compile(
    r"\\(?:biggl|biggr|Biggl|Biggr|bigg|Bigg|bigl|bigr|Bigl|Bigr|big|Big)\b\s*"
)


def _get_xslt():
    """Load and cache the XSLT stylesheet."""
    global _xslt
    if _xslt is None:
        _xslt = etree.parse(str(_XSL_PATH))
    return _xslt


def normalize_latex(latex_str: str) -> str:
    """Normalize OCR LaTeX before preview/conversion."""
    latex = latex_str.strip()
    wrappers = (
        ("\\[", "\\]"),
        ("\\(", "\\)"),
        ("$$", "$$"),
        ("$", "$"),
    )
    for start, end in wrappers:
        if latex.startswith(start) and latex.endswith(end):
            latex = latex[len(start):-len(end)].strip()
            break
    latex = _normalize_script_arguments(latex)
    return latex


def _normalize_script_arguments(latex: str) -> str:
    """Add braces around unbraced superscript/subscript arguments.

    OCR engines often produce valid-looking but ambiguous output such as
    ``x^ij`` or ``a_n+1``. This keeps already-braced input unchanged and only
    groups a conservative next-token argument.
    """
    result = []
    i = 0
    while i < len(latex):
        char = latex[i]
        if char not in "^_":
            result.append(char)
            i += 1
            continue

        result.append(char)
        i += 1
        while i < len(latex) and latex[i].isspace():
            i += 1

        if i >= len(latex):
            break

        if latex[i] == "{":
            group, i = _consume_balanced_group(latex, i)
            result.append(group)
            continue

        if latex[i] in "+-" and i + 1 < len(latex):
            start = i
            i += 1
            if latex[i] == "\\":
                command_match = re.match(r"\\[A-Za-z]+|\\.", latex[i:])
                if command_match:
                    i += len(command_match.group(0))
                    result.append(f"{{{latex[start:i]}}}")
                    continue
            while i < len(latex) and latex[i].isalnum():
                i += 1
            result.append(f"{{{latex[start:i]}}}")
            continue

        if latex[i] == "\\":
            command_match = re.match(r"\\[A-Za-z]+|\\.", latex[i:])
            if command_match:
                arg = command_match.group(0)
                i += len(arg)
                result.append(f"{{{arg}}}")
                continue

        if latex[i].isalnum():
            start = i
            while i < len(latex) and latex[i].isalnum():
                i += 1
            result.append(f"{{{latex[start:i]}}}")
            continue

        result.append(f"{{{latex[i]}}}")
        i += 1

    return "".join(result)


def _consume_balanced_group(text: str, start: int) -> tuple[str, int]:
    depth = 0
    i = start
    while i < len(text):
        if text[i] == "{":
            depth += 1
        elif text[i] == "}":
            depth -= 1
            if depth == 0:
                return text[start:i + 1], i + 1
        i += 1
    return text[start:], len(text)


def _to_matplotlib_mathtext(latex_str: str) -> str:
    """Return a preview-only LaTeX variant accepted by Matplotlib mathtext."""
    latex = normalize_latex(latex_str)
    return _MATHTEXT_SIZE_COMMAND_RE.sub("", latex)


def _to_matplotlib_plain_text(latex_str: str) -> str:
    """Return a literal text fallback that will not trigger mathtext parsing."""
    return normalize_latex(latex_str).replace("$", r"\$")


def _xml_to_string(element: etree._Element) -> str:
    """Serialize an XML element with stable pretty formatting."""
    return etree.tostring(
        element,
        encoding="unicode",
        pretty_print=True,
    )


def latex_to_mathml(latex_str: str) -> str:
    """Convert LaTeX string to MathML XML string."""
    return latex2mathml.converter.convert(normalize_latex(latex_str))


def mathml_to_omml(mathml_str: str) -> etree._Element:
    """Convert MathML XML string to OMML XML element using Microsoft's XSLT."""
    xslt = _get_xslt()
    mathml_tree = etree.fromstring(mathml_str.encode("utf-8"))
    transform = etree.XSLT(xslt)
    omml_tree = transform(mathml_tree)
    return omml_tree.getroot()


def latex_to_omml(latex_str: str) -> etree._Element:
    """Convert LaTeX string directly to OMML XML element."""
    mathml = latex_to_mathml(latex_str)
    return mathml_to_omml(mathml)


def latex_to_omml_string(latex_str: str) -> str:
    """Convert LaTeX string to Word Office MathML (OMML) XML."""
    return _xml_to_string(latex_to_omml(latex_str))


def create_word_doc(latex_str: str, output_path: str) -> None:
    """Create a Word document with a native editable equation.

    Args:
        latex_str: LaTeX formula string.
        output_path: Path to save the .docx file.
    """
    doc = Document()

    # Add title
    title = doc.add_paragraph("Formula")
    title.style = doc.styles["Title"]

    # Add equation paragraph
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Convert LaTeX to OMML
    omml_element = latex_to_omml(latex_str)

    # Inject OMML into the paragraph
    para._p.append(omml_element)

    # Also add the raw LaTeX as a reference paragraph
    ref = doc.add_paragraph()
    ref.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = ref.add_run(f"LaTeX: {normalize_latex(latex_str)}")
    run.font.size = 10
    run.font.color.rgb = None  # Default color

    doc.save(output_path)


def create_word_doc_multi(latex_list: list[str], output_path: str) -> None:
    """Create a Word document with multiple equations.

    Args:
        latex_list: List of LaTeX formula strings.
        output_path: Path to save the .docx file.
    """
    doc = Document()
    doc.add_paragraph("Formulas").style = doc.styles["Title"]

    for i, latex_str in enumerate(latex_list, 1):
        # Equation number
        heading = doc.add_paragraph(f"Equation {i}")
        heading.style = doc.styles["Heading 2"]

        # Native equation
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        omml_element = latex_to_omml(latex_str)
        para._p.append(omml_element)

        # LaTeX source
        ref = doc.add_paragraph()
        ref.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = ref.add_run(f"LaTeX: {latex_str}")
        run.font.size = 10

        doc.add_paragraph()  # Spacing

    doc.save(output_path)


def _draw_preview_image(
    text: str,
    dpi: int,
    *,
    math_mode: bool,
    fontsize: int,
) -> Image.Image:
    import matplotlib
    matplotlib.use("Agg")
    import matplotlib.pyplot as plt
    from matplotlib.figure import Figure

    fig = Figure(figsize=(6, 1.5), dpi=dpi)
    ax = fig.add_axes([0, 0, 1, 1])
    ax.set_axis_off()
    display_text = text
    if math_mode and not display_text.startswith("$"):
        display_text = f"${display_text}$"

    ax.text(
        0.5, 0.5, display_text,
        fontsize=fontsize,
        ha="center", va="center",
        transform=ax.transAxes,
        family=None if math_mode else "monospace",
    )

    buf = BytesIO()
    try:
        fig.savefig(buf, format="png", dpi=dpi, bbox_inches="tight", pad_inches=0.1)
    finally:
        plt.close(fig)
    buf.seek(0)
    return Image.open(buf)


def render_latex_preview(latex_str: str, dpi: int = 150) -> Image.Image:
    """Render LaTeX formula to a PIL Image for preview.

    Args:
        latex_str: LaTeX formula string.
        dpi: Resolution for rendering.

    Returns:
        PIL Image of the rendered formula.
    """
    preview_tex = _to_matplotlib_mathtext(latex_str)
    try:
        return _draw_preview_image(preview_tex, dpi, math_mode=True, fontsize=20)
    except Exception:
        fallback_text = _to_matplotlib_plain_text(latex_str)
        return _draw_preview_image(fallback_text, dpi, math_mode=False, fontsize=14)
